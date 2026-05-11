import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill, color=None, val="clear"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcShd = OxmlElement('w:shd')
    tcShd.set(qn('w:fill'), fill)
    tcShd.set(qn('w:val'), val)
    if color:
        tcShd.set(qn('w:color'), color)
    tcPr.append(tcShd)

def add_border(cell, **kwargs):
    """
    kwargs can contain: top, bottom, start, end
    each being a dictionary with attributes like: val, color, sz, space
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in kwargs[edge].items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

data = [
    ["1", "TRAVESSA CAMBARATIBA", "-", "260,00", "Pendente"],
    ["2", "CLEONICE KAMMER DI SANDRO", "-", "1.405,00", "Pendente"],
    ["3", "ANA MARIA FRANCO LARANJEIRA", "-", "785,00", "Pendente"],
    ["4", "RUA ILHA DA VITORIA", "Trecho Estrada São Paulo - Jundiai a Escola do Badra", "662,00", "Pendente"],
    ["5", "ILHA DO FRADE", "Trecho parcial com inicio na Ilha da Vitoria", "250,00", "Concluído"],
    ["6", "RUA MAJORLANDIA", "Trecho Estrada da ligação a Rua Artur de Azevedo", "351,00", "Pendente"],
    ["7", "ESTRADA DE PIRAPORA", "(Antiga João Fernandes Vieira) ao seu final", "205,00", "Pendente"],
    ["8", "NOEL ROSA", "Trecho Rua do Educador a Antonio Conselheiro", "242,00", "Pendente"],
    ["9", "PADRE JOSINO", "-", "445,00", "Concluído"],
    ["10", "FORTALEZA", "-", "548,00", "Concluído"],
    ["11", "SANTO DIAS", "-", "428,00", "Concluído"],
    ["12", "BETINHO", "-", "395,00", "Concluído"],
    ["13", "RUA MARILIA", "Trecho final", "243,00", "Pendente"],
    ["14", "LEONEL MARTINIANO", "Chácara Maria Trindade", "1.565,00", "Pendente"],
    ["15", "MOMBAÇA DE CIMA", "Morro da Mandioca", "560,00", "Pendente"],
    ["16", "RUA DO ANCIÃO", "Morro da Mandioca", "268,00", "Pendente"],
    ["17", "MANOEL JOAQUIM DE SANT'ANA", "Morro da Mandioca | Jardim Jaraguá", "1.333,00", "Pendente"],
    ["18", "RUA BELA VISTA JD DA PAZ", "Fanton | Sehab está com obra no local, não se sabe se farão pavimentação.", "160,00", "Concluído"],
    ["19", "RUA NOVA VIDA JD DAPAZ", "Fanton | Sehab está com obra no local, não se sabe se farão pavimentação.", "170,00", "Concluído"],
    ["20", "RUA ESPERANÇA JD DA PAZ", "Fanton | Sehab está com obra no local, não se sabe se farão pavimentação.", "160,00", "Concluído"],
    ["21", "RUA MORAIS JD DA PAZ", "Fanton | Sehab está com obra no local, não se sabe se farão pavimentação.", "280,00", "Concluído"],
    ["22", "RUA NOVA CANAÃ JD DA PAZ", "Fanton | Sehab está com obra no local, não se sabe se farão pavimentação.", "228,00", "Concluído"],
    ["23", "RUA VIEIRA DE BRITO BOTUQUARA", "(Bloco de Rocha)", "140,00", "Pendente"],
    ["24", "RUA ARTUR DE AZEVEDO E RUA DE SERVIDÃO", "Sol Nascente", "-", "Pendente"]
]

# Create document
doc = Document()

# Set Landscape orientation
sections = doc.sections
for section in sections:
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
title = doc.add_heading('RELAÇÃO DE OBRAS A SEREM LEVANTADAS PARA ORÇAR 2024', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x0F, 0x20, 0x4B) # Ultra dark premium blue
    run.font.bold = True
    run.font.size = Pt(18)

# Header Date info
date_para = doc.add_paragraph('ENVIADOS PELA LUCIANA EM 09.01.2024')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
d_run = date_para.runs[0]
d_run.font.color.rgb = RGBColor(0x90, 0x00, 0x00) # Deep red for date
d_run.font.bold = True
d_run.font.size = Pt(11)

doc.add_paragraph()

# Add the 4 intro notes as styled bullet points
notes = [
    "Rua Chamburcy (dá acesso para área municipal)",
    "Rua Ana Maria Franco Laranjeiras",
    "Rua Cleonice Kammer Di Sandro",
    "Alameda Aristóteles Claudio Sbrig (Parte é recape e parte é pavimentação nova)"
]
notes_p = doc.add_paragraph()
notes_run = notes_p.add_run("Vias em Destaque:\n")
notes_run.font.bold = True
notes_run.font.color.rgb = RGBColor(0x0F, 0x20, 0x4B)
for note in notes:
    p = doc.add_paragraph(f"• {note}")
    p.paragraph_format.left_indent = Inches(0.2)
    p.runs[0].font.size = Pt(10)

doc.add_paragraph() # Spacing

# Add table
table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Set column widths (Total ~ 9.4 inches)
widths = [Inches(0.6), Inches(3.0), Inches(3.8), Inches(1.0), Inches(1.0)]
for i, col in enumerate(table.columns):
    for cell in col.cells:
        cell.width = widths[i]

# Header
hdr_cells = table.rows[0].cells
headers = ['ITEM', 'NOME DO LOGRADOURO', 'TRECHO / OBSERVAÇÕES', 'EXTENSÃO (m)', 'STATUS']

for i, header in enumerate(headers):
    hdr_cells[i].text = header
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text
    set_cell_background(hdr_cells[i], "0F204B") # Premium dark blue
    # Add bottom border
    add_border(hdr_cells[i], bottom={'val': 'single', 'color': '0F204B', 'sz': '12'})

# Add Data
for row_idx, row_data in enumerate(data):
    row_cells = table.add_row().cells
    
    # Subtle alternation
    bg_color = "F8F9FA" if row_idx % 2 == 0 else "FFFFFF"
    
    for i, val in enumerate(row_data):
        row_cells[i].text = str(val)
        p = row_cells[i].paragraphs[0]
        
        # Background
        set_cell_background(row_cells[i], bg_color)
        
        # Add thin borders
        add_border(row_cells[i], 
                   bottom={'val': 'single', 'color': 'E0E0E0', 'sz': '4'},
                   top={'val': 'single', 'color': 'E0E0E0', 'sz': '4'})
        
        # Alignment
        if i in (0, 3, 4): # Item, Extensao, Status
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Font styling
        run = p.runs[0]
        run.font.size = Pt(10)
        
        if i == 1: # Logradouro
            run.font.bold = True
        
        if i == 2: # Observacoes
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55) # subtle grey
            run.font.italic = True
        
        # Color coding for status
        if i == 4:
            if val == 'Pendente':
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) # Dark Red
                run.font.bold = True
            elif val == 'Concluído':
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Dark Green
                run.font.bold = True

# Add Total Row
row_cells = table.add_row().cells
for i in range(5):
    set_cell_background(row_cells[i], "0F204B")
    
row_cells[1].text = "TOTAL GERAL"
p = row_cells[1].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.runs[0]
run.font.bold = True
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

row_cells[3].text = "11.083,00"
p = row_cells[3].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]
run.font.bold = True
run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.save('Relatorio de Obras 2024 - Elite.docx')
print("Documento salvo como 'Relatorio de Obras 2024 - Elite.docx'")

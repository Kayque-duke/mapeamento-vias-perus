import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

data = [
    ["Perus", "Travessa Cambá", "5.00", "Pendente"],
    ["Sítio Areião", "Rua Ana Maria F", "8.00", "Pendente"],
    ["Sítio Areião", "Rua Cleonice Ka", "8.00", "Pendente"],
    ["Sítio Botuquara", "Rua Vieira de Br", "5.00", "Pendente"],
    ["Jd. Nascente", "Rua Artur de Azi", "8.00", "Pendente"],
    ["Jd. Nascente", "Alameda Aristót", "5.50", "Pendente"],
    ["Vila Harlon", "Rua Esperança", "5.00", "Concluído"],
    ["Vila Palmares", "Estrada do Pirap", "6.50", "Concluído"],
    ["Maria Trindade", "Rua Mar do Nort", "5.00", "Pendente"],
    ["Maria Trindade", "Rua Mar das Flo", "5.00", "Pendente"],
    ["Maria Trindade", "Rua mar da Irlan", "5.00", "Pendente"],
    ["J. Jaraguá", "Rua do Marajó", "10.71", "Pendente"],
    ["Maria Trindade", "Rua Ilha de Bali", "8.13", "Pendente"],
    ["Jardim Jaraguá", "Rua Ilhas Britâni", "8.52", "Pendente"],
    ["Maria Trindade", "Rua Guaratinga", "6.00", "Pendente"],
    ["Maria Trindade", "Rua Sapucainha", "6.00", "Pendente"],
    ["Maria Trindade", "Rua das Alpinas", "6.00", "Pendente"],
    ["M. Trindade", "Rua do Cerrado", "6.00", "Pendente"],
    ["Riscado/Vila Per", "Rua Ilha da Vitor", "6.00", "Concluído"],
]

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('Levantamento de Vias Sem Pavimentação', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A) # Modern blue

# Subtitle
subtitle = doc.add_paragraph('Relatório de Visitas Pendentes - Perus 2024')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.bold = True
subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph() # Spacing

# Add table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid' # Standard Word styling
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Set column widths
widths = [Inches(1.5), Inches(2.5), Inches(1.0), Inches(1.5)]
for i, col in enumerate(table.columns):
    for cell in col.cells:
        cell.width = widths[i]

# Header
hdr_cells = table.rows[0].cells
headers = ['Bairro/Subprefeitura', 'Logradouro', 'Largura (m)', 'Status da Foto']
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add Data
for row_data in data:
    # Filter? No, I'll keep all of them but color code the status
    row_cells = table.add_row().cells
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        p = row_cells[i].paragraphs[0]
        # Align center for Largura and Status
        if i >= 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Color coding for status
        if i == 3:
            if val == 'Pendente':
                p.runs[0].font.color.rgb = RGBColor(0xD9, 0x53, 0x4F) # Red
                p.runs[0].font.bold = True
            elif val == 'Concluído':
                p.runs[0].font.color.rgb = RGBColor(0x5C, 0xB8, 0x5C) # Green
                p.runs[0].font.bold = True

doc.add_paragraph()
footer = doc.add_paragraph('Documento gerado automaticamente pelo sistema Genius Word Master.')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(10)
footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save('Visitas pendentes.docx')
print("Documento salvo como 'Visitas pendentes.docx'")

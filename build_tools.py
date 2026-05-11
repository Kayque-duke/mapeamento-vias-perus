import json
import os
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- DATA PREPARATION ---
route_centers = {
    "Rota 1": {"lat": -23.4115, "lng": -46.8135},
    "Rota 2": {"lat": -23.4466, "lng": -46.7892},
    "Rota 3": {"lat": -23.4135, "lng": -46.7379},
    "Rota 4": {"lat": -23.4201, "lng": -46.7196},
    "Rota 5": {"lat": -23.4037, "lng": -46.7539}
}

with open('rotas_coordenadas.json', 'r', encoding='utf-8') as f:
    routes_data = json.load(f)

# Normalize missing coordinates and add some slight jitter so they don't overlap exactly
processed_streets = []
import random
random.seed(42)

for route, vias in routes_data.items():
    center = route_centers[route]
    for i, via in enumerate(vias):
        lat = via['lat']
        lng = via['lng']
        
        # If null or way off (e.g. -22 or -23.56 which are wrong cities/downtown), use center + jitter
        if lat is None or lat > -23.3 or lat < -23.45:
            lat = center['lat'] + random.uniform(-0.002, 0.002)
            lng = center['lng'] + random.uniform(-0.002, 0.002)
            
        processed_streets.append({
            "id": f"{route}_{i}",
            "name": via['name'],
            "width": 6,
            "length": 150,
            "start": {"lat": lat, "lng": lng},
            "end": {"lat": lat + 0.0005, "lng": lng + 0.0005},
            "type": "Rua",
            "pavement": "Asfalto",
            "obs": route
        })

# --- GENERATE HTML MAP ---
with open('genio_ruas_perus.html', 'r', encoding='utf-8') as f:
    html_content = f.read()

# Inject the processed streets into the html
# We will find the var streets = []; or similar initialization.
injection_script = f"""
// INJECTED BY PLANNER
let preloaded_streets = {json.dumps(processed_streets)};
if (typeof streets !== 'undefined') {{
    streets = preloaded_streets;
}} else {{
    window.streets = preloaded_streets;
}}
// END INJECTION
"""

html_out = html_content.replace('let streets = [];', f'let streets = {json.dumps(processed_streets)};')
if 'let streets =' not in html_out:
    # try var
    html_out = html_out.replace('var streets = [];', f'var streets = {json.dumps(processed_streets)};')

# Also set map center to Perus
html_out = html_out.replace('[-23.5505, -46.6333]', '[-23.407, -46.756]') # SP center to Perus center
html_out = html_out.replace('<title>CPO - Mapeamento Viário Perus</title>', '<title>CPO - Mapeamento Especial de Rotas</title>')

with open('genio_rotas_vistoria.html', 'w', encoding='utf-8') as f:
    f.write(html_out)


# --- GENERATE WORD DOCUMENT ---
doc = Document()

# Add styles
style = doc.styles['Normal']
font = style.font
font.name = 'Inter'
font.size = Pt(11)

# Title
title = doc.add_heading('Plano de Vistorias - Perus e Jaraguá', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.color.rgb = RGBColor(0x0F, 0x11, 0x17)

doc.add_paragraph("Este documento organiza as 20 vias solicitadas em 5 rotas otimizadas por proximidade geográfica, projetado para eficiência em campo.", style='Normal')

# Organize content
route_descriptions = {
    "Rota 1": "Loteamento Maria Trindade - Região com a maior concentração de ruas pendentes.",
    "Rota 2": "Jardim Jaraguá - Ruas localizadas no setor do Jaraguá formando a segunda rota.",
    "Rota 3": "Jd. Nascente & Sítio Areião - Loteamentos vizinhos conectados no mesmo trajeto.",
    "Rota 4": "Botuquara - Afastadas do centro, mas próximas entre si.",
    "Rota 5": "Área Central (Perus) & Morro Doce - Vias fáceis de encaixar no final ou início do dia."
}

for route in sorted(routes_data.keys()):
    h = doc.add_heading(route, level=1)
    h.runs[0].font.color.rgb = RGBColor(0xE3, 0x06, 0x13) # Vermelho SP
    
    doc.add_paragraph(route_descriptions.get(route, ""), style='Normal')
    
    # Table
    vias = routes_data[route]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Status'
    hdr_cells[1].text = 'Nome da Via'
    hdr_cells[2].text = 'Observações / Anotações'
    
    for via in vias:
        row_cells = table.add_row().cells
        row_cells[0].text = '[   ]'
        row_cells[1].text = via['name']
        row_cells[2].text = ''
        
    doc.add_paragraph() # space

doc.save('Plano_Vistorias_Perus_Jaragua.docx')
print("Artifacts generated successfully.")

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

data = [
    ["Perus", "Travessa Cambaratiba", "Sim", "Sim", "5,00", "Pendente"],
    ["Sítio Areião", "Rua Ana Maria Franco Laranjeiras", "Sim", "Sim", "8,00", "Pendente"],
    ["Sítio Areião", "Rua Cleonice Kammer di Sandro", "Sim", "Sim", "8,00", "Pendente"],
    ["Sítio Botuquara", "Rua Vieira de Brito", "Sim", "Sim", "5,00", "Pendente"],
    ["Jd. Nascente", "Rua Artur de Azevedo", "Sim", "Sim", "8,00", "Pendente"],
    ["Jd. Nascente", "Alameda Aristoteles Claudio Sbrigh", "Sim", "Sim", "5,50", "Pendente"],
    ["Vila Harlon", "Rua Esperança", "Sim", "Não", "5,00", "Concluído"],
    ["Vila Palmares", "Estrada do Pirapora (feita parcial)", "Sim", "Sim", "6,50", "Concluído"],
    ["Maria Trindade", "Rua Mar do Norte", "Sim", "Sim", "5,00", "Pendente"],
    ["Maria Trindade", "Rua Mar das Flores", "Sim", "Sim", "5,00", "Pendente"],
    ["Maria Trindade", "Rua mar da Irlanda", "Sim", "Sim", "5,00", "Pendente"],
    ["J. Jaraguá", "Rua do Marajó", "Sim", "Sim", "10,71", "Pendente"],
    ["Maria Trindade", "Rua Ilha de Bali", "Sim", "Sim", "8,13", "Pendente"],
    ["Jardim Jaraguá", "Rua Ilhas Britânicas", "Sim", "Sim", "8,52", "Pendente"],
    ["Maria Trindade", "Rua Guaratinga", "Sim", "Não", "6,00", "Pendente"],
    ["Maria Trindade", "Rua Sapucainha", "Sim", "Não", "6,00", "Pendente"],
    ["Maria Trindade", "Rua das Alpinas", "Sim", "Não", "6,00", "Pendente"],
    ["Maria Trindade", "Rua do Cerrado Brasileiro", "Sim", "Não", "6,00", "Pendente"],
    ["Vila Perus", "Rua Ilha da Vitoria", "Não", "Sim", "6,00", "Concluído"],
    ["Vila Perus", "Rua Ilha do Frade", "Sim", "Sim", "6,00", "Pendente"],
    ["Vila Perus", "Rua Chamburcy", "Sim", "Sim", "6,00", "Pendente"],
    ["Botuquara", "Rua Miguel Vilela", "Sim", "Sim", "6,00", "Pendente"],
    ["Vila Palmares", "Rua Santo Dias", "Sim", "Não", "6,00", "Concluído"],
    ["Jardim Jaraguá", "Rua Denis Brean", "Sim", "Sim", "6,00", "Concluído"],
    ["Jardim Jaraguá", "Rua Albert Jansen", "Não", "Sim", "5,00", "Pendente"],
    ["-", "Travessa Lobito", "Sim", "Sim", "5,00", "Pendente"],
    ["Vila Palmares", "Rua Betinho", "Sim", "Não", "5,00", "Concluído"],
    ["S. Jaraguá", "VL onze", "Sim", "Não", "6,00", "Concluído"],
    ["Vila Palmares", "Rua Fortaleza", "Sim", "Não", "6,00", "Concluído"],
]

# Create document
doc = Document()

# Set Landscape orientation
sections = doc.sections
for section in sections:
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
title = doc.add_heading('Relação de Obras a Serem Levantadas - Perus 2024', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Deep blue

# Subtitle
subtitle = doc.add_paragraph('Levantamento Completo de Vias Sem Pavimentação')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.bold = True
subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph() # Spacing

# Add table
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Set column widths (Total ~ 9.4 inches)
widths = [Inches(1.5), Inches(3.5), Inches(0.9), Inches(1.1), Inches(1.0), Inches(1.4)]
for i, col in enumerate(table.columns):
    for cell in col.cells:
        cell.width = widths[i]

# Header
hdr_cells = table.rows[0].cells
headers = ['Bairro / Subprefeitura', 'Nome do Logradouro', 'Via\nInteira', 'Logradouro\nOficial', 'Largura (m)', 'Status']

# Style Header
import docx.oxml.shared
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill, color=None, val="clear"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcShd = OxmlElement('w:shd')
    tcShd.set(qn('w:fill'), fill)
    tcShd.set(qn('w:val'), val)
    if color:
        tcShd.set(qn('w:color'), color)
    tcPr.append(tcShd)

for i, header in enumerate(headers):
    hdr_cells[i].text = header
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text
    set_cell_background(hdr_cells[i], "1F4E79") # Deep blue background

# Add Data
for row_idx, row_data in enumerate(data):
    row_cells = table.add_row().cells
    
    # Alternate row colors for "fenomenal" look
    bg_color = "F2F2F2" if row_idx % 2 == 0 else "FFFFFF"
    
    for i, val in enumerate(row_data):
        row_cells[i].text = val
        p = row_cells[i].paragraphs[0]
        
        # Background
        set_cell_background(row_cells[i], bg_color)
        
        # Alignment
        if i >= 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Font styling
        run = p.runs[0]
        run.font.size = Pt(10)
        
        # Color coding for status
        if i == 5:
            if val == 'Pendente':
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) # Dark Red
                run.font.bold = True
            elif val == 'Concluído':
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Dark Green
                run.font.bold = True

doc.add_paragraph()
footer = doc.add_paragraph('Documento estruturado e automatizado sob as diretrizes do Genius Word Master.')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
footer.runs[0].font.italic = True

doc.save('Visitas Pendentes - Atualizado.docx')
print("Documento salvo como 'Visitas Pendentes - Atualizado.docx'")

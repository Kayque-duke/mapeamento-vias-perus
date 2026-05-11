import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill, color=None, val="clear"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcShd = OxmlElement('w:shd')
    tcShd.set(qn('w:fill'), fill)
    tcShd.set(qn('w:val'), val)
    if color:
        tcShd.set(qn('w:color'), color)
    tcPr.append(tcShd)

def add_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in kwargs[edge].items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

data = [
    ["Travessa Cambaratiba", "Perus", "Pendente"],
    ["Rua Cleonice Kammer Di Sandro", "Sítio Areião", "Pendente"],
    ["Rua Ana Maria Franco Laranjeira", "Sítio Areião", "Pendente"],
    ["Rua Ilha da Vitoria", "Vila Perus", "Pendente"],
    ["Rua Ilha do Frade", "Vila Perus", "Concluído"],
    ["Rua Majorlandia", "Não Especificado", "Pendente"],
    ["Estrada de Pirapora", "Vila Palmares", "Pendente"],
    ["Rua Noel Rosa", "Não Especificado", "Pendente"],
    ["Rua Padre Josino", "Não Especificado", "Concluído"],
    ["Rua Fortaleza", "Vila Palmares", "Concluído"],
    ["Rua Santo Dias", "Vila Palmares", "Concluído"],
    ["Rua Betinho", "Vila Palmares", "Concluído"],
    ["Rua Marilia", "Não Especificado", "Pendente"],
    ["Rua Leonel Martiniano", "Maria Trindade", "Pendente"],
    ["Rua Mombaça de Cima", "Morro da Mandioca", "Pendente"],
    ["Rua do Ancião", "Morro da Mandioca", "Pendente"],
    ["Rua Manoel Joaquim de Sant'ana", "Jardim Jaraguá", "Pendente"],
    ["Rua Bela Vista", "Jd. da Paz", "Concluído"],
    ["Rua Nova Vida", "Jd. da Paz", "Concluído"],
    ["Rua Esperança", "Jd. da Paz", "Concluído"],
    ["Rua Morais", "Jd. da Paz", "Concluído"],
    ["Rua Nova Canaã", "Jd. da Paz", "Concluído"],
    ["Rua Vieira de Brito", "Botuquara", "Pendente"],
    ["Rua Artur de Azevedo e Servidão", "Sol Nascente", "Pendente"],
    ["Rua Chamburcy", "Vila Perus", "Pendente"],
    ["Alameda Aristoteles Claudio Sbrig", "Jd. Nascente", "Pendente"]
]

# Create document
doc = Document()

# Set Margins (Portrait mode is fine for 3 columns)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('Relatório de Vias e Status', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x0F, 0x20, 0x4B) # Ultra dark premium blue
    run.font.bold = True

doc.add_paragraph() # Spacing

# Add table
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# Set column widths (Total ~ 6.5 inches)
widths = [Inches(3.0), Inches(2.0), Inches(1.5)]
for i, col in enumerate(table.columns):
    for cell in col.cells:
        cell.width = widths[i]

# Header
hdr_cells = table.rows[0].cells
headers = ['NOME DA RUA', 'BAIRRO', 'STATUS']

for i, header in enumerate(headers):
    hdr_cells[i].text = header
    p = hdr_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text
    set_cell_background(hdr_cells[i], "0F204B") # Premium dark blue
    add_border(hdr_cells[i], bottom={'val': 'single', 'color': '0F204B', 'sz': '12'})

# Add Data
for row_idx, row_data in enumerate(data):
    row_cells = table.add_row().cells
    
    # Subtle alternation
    bg_color = "F8F9FA" if row_idx % 2 == 0 else "FFFFFF"
    
    for i, val in enumerate(row_data):
        row_cells[i].text = str(val)
        p = row_cells[i].paragraphs[0]
        
        # Background
        set_cell_background(row_cells[i], bg_color)
        
        # Add thin borders
        add_border(row_cells[i], 
                   bottom={'val': 'single', 'color': 'E0E0E0', 'sz': '4'},
                   top={'val': 'single', 'color': 'E0E0E0', 'sz': '4'},
                   left={'val': 'single', 'color': 'E0E0E0', 'sz': '4'},
                   right={'val': 'single', 'color': 'E0E0E0', 'sz': '4'})
        
        # Alignment
        if i in (1, 2): # Bairro and Status
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Font styling
        run = p.runs[0]
        run.font.size = Pt(11)
        
        if i == 0: # Nome da rua
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            
        if i == 1: # Bairro
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        # Color coding for status
        if i == 2:
            if val == 'Pendente':
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) # Dark Red
                run.font.bold = True
            elif val == 'Concluído':
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Dark Green
                run.font.bold = True

doc.save('Status das Ruas - Limpo e Direto.docx')
print("Documento salvo como 'Status das Ruas - Limpo e Direto.docx'")
